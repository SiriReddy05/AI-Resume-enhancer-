from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import PyPDF2
import docx
import re
import json
import os
from groq import Groq
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors

# Add these new imports for DOCX generation
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
# Allow CORS for all domains so Vercel can access it!
CORS(app, resources={r"/*": {"origins": "*"}})

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

@app.route('/')
def home():
    return "Backend is running!"

def extract_pdf(file_path):
    text = ""
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def extract_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# 🧠 AI Enhancement (Strict JSON Structured Output)
def enhance_with_ai(raw_text):
    try:
        system_prompt = """
        You are an expert ATS-friendly resume writer and a meticulous data extractor.
        
        CRITICAL RULES:
        1. YOU ARE FORBIDDEN FROM DELETING OR SUMMARIZING INFORMATION.
        2. You MUST retain EVERY single job experience, EVERY project, EVERY certification, EVERY achievement, and EVERY bullet point.
        3. EXPLICIT EXTRACTION: Actively search for "Scholastic Achievements", "Position of Responsibility", or "Leadership". Extract ALL bullet points into the "achievements" array.
        4. CATEGORIZED SKILLS: Group skills into categories (e.g., Languages, Frameworks).
        5. EXTRACT TECHNOLOGIES: For every job experience and project, explicitly extract the technologies/skills used into the "technologies" field.
        
        You MUST return ONLY a valid JSON object. Do not include markdown code blocks (no ```json).
        Use this EXACT JSON structure:
        {
          "name": "Full Name",
          "contact": "Email | Phone | LinkedIn/Links",
          "summary": "A highly professional summary paragraph...",
          "skills": [
            {"category": "Languages", "items": "Python, JavaScript, SQL"}
          ],
          "experience": [
            {
              "title": "Job Title or Role",
              "company": "Company Name",
              "technologies": "Python, OpenAI, Azure AI, Docker",
              "location": "City, Country",
              "date": "Duration (e.g., Jan 2021 - Present)",
              "details": ["High-impact bullet point 1"]
            }
          ],
          "projects": [
            {
              "title": "Project Name",
              "technologies": "ReactJS, Node.js, Python",
              "date": "Year or Duration",
              "details": ["Action-oriented bullet 1"]
            }
          ],
          "education": [
            {
              "title": "Degree Name",
              "subtitle": "Institution Name",
              "date": "Graduation Year",
              "details": ["Optional academic detail or empty string"]
            }
          ],
          "certifications": ["Certification 1"],
          "achievements": ["Achievement/Leadership 1"]
        }
        """

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"RAW RESUME TEXT:\n{raw_text}"}
            ],
            temperature=0.1,
            max_tokens=6000
        )

        response_text = response.choices[0].message.content.strip()
        
        if response_text.startswith("```json"): 
            response_text = response_text[7:]
        if response_text.startswith("```"): 
            response_text = response_text[3:]
        if response_text.endswith("```"): 
            response_text = response_text[:-3]

        return json.loads(response_text.strip())

    except Exception as e:
        print("AI ERROR:", str(e))
        return {"error": "AI failed to process the resume properly."}
    
def generate_pdf(data, filename="improved_resume.pdf"):
    doc = SimpleDocTemplate(filename, pagesize=letter, rightMargin=25, leftMargin=25, topMargin=20, bottomMargin=20)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(name='TitleStyle', alignment=TA_CENTER, fontSize=15, leading=18, spaceAfter=2, fontName='Helvetica-Bold')
    contact_style = ParagraphStyle(name='ContactStyle', alignment=TA_CENTER, fontSize=9, spaceAfter=8)
    section_heading = ParagraphStyle(name='SectionHeading', fontSize=10, spaceBefore=6, spaceAfter=2, textTransform='uppercase', fontName='Helvetica-Bold')
    item_title = ParagraphStyle(name='ItemTitle', fontSize=9, spaceBefore=3, spaceAfter=0, fontName='Helvetica-Bold')
    normal_text = ParagraphStyle(name='NormalText', fontSize=9, leading=10, spaceAfter=2)
    bullet_text = ParagraphStyle(name='BulletText', fontSize=9, leading=10, leftIndent=10, firstLineIndent=0, spaceAfter=1)

    elements = []

    elements.append(Paragraph(data.get('name', '').upper(), title_style))
    elements.append(Paragraph(data.get('contact', ''), contact_style))
    
    def add_line():
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceBefore=0, spaceAfter=2))

    if data.get('summary'):
        elements.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", section_heading))
        add_line()
        elements.append(Paragraph(data['summary'], normal_text))

    if data.get('skills'):
        elements.append(Paragraph("<b>SKILLS</b>", section_heading))
        add_line()
        for skill_group in data['skills']:
            cat = skill_group.get('category', '')
            items = skill_group.get('items', '')
            if cat and items:
                elements.append(Paragraph(f"<b>{cat}:</b> {items}", normal_text))
            elif items:
                elements.append(Paragraph(f"• {items}", normal_text))

    # --- UPDATED EXPERIENCE LAYOUT ---
    if data.get('experience'):
        elements.append(Paragraph("<b>EXPERIENCE</b>", section_heading))
        add_line()
        for item in data['experience']:
            header = f"<b>{item.get('title', '')}</b>"
            if item.get('date'): header += f" <font color='grey'>| {item.get('date', '')}</font>"
            elements.append(Paragraph(header, item_title))
            
            sub_parts = []
            if item.get('company'): sub_parts.append(f"<b>{item['company']}</b>")
            if item.get('technologies'): sub_parts.append(item['technologies'])
            sub_str = " | ".join(sub_parts)
            if item.get('location'): sub_str += f" <font color='grey'>| {item['location']}</font>"
            
            if sub_str: elements.append(Paragraph(f"<i>{sub_str}</i>", normal_text))
            for detail in item.get('details', []):
                if detail.strip(): elements.append(Paragraph(f"• {detail}", bullet_text))

    # --- UPDATED PROJECTS LAYOUT ---
    if data.get('projects'):
        elements.append(Paragraph("<b>PROJECTS</b>", section_heading))
        add_line()
        for item in data['projects']:
            header = f"<b>{item.get('title', '')}</b>"
            if item.get('technologies'): header += f" | <i><font color='dimgrey'>{item['technologies']}</font></i>"
            if item.get('date'): header += f" <font color='grey'>| {item.get('date', '')}</font>"
            elements.append(Paragraph(header, item_title))
            
            for detail in item.get('details', []):
                if detail.strip(): elements.append(Paragraph(f"• {detail}", bullet_text))

    if data.get('education'):
        elements.append(Paragraph("<b>EDUCATION</b>", section_heading))
        add_line()
        for edu in data['education']:
            elements.append(Paragraph(f"<b>{edu.get('title', '')}</b> | {edu.get('date', '')}", item_title))
            elements.append(Paragraph(edu.get('subtitle', ''), normal_text))

    for section_title, key in [("CERTIFICATIONS", "certifications"), ("ACHIEVEMENTS & LEADERSHIP", "achievements")]:
        items = data.get(key, [])
        if items:
            elements.append(Paragraph(f"<b>{section_title}</b>", section_heading))
            add_line()
            for item in items:
                if str(item).strip(): elements.append(Paragraph(f"• {item}", bullet_text))

    doc.build(elements)
    return filename

def generate_docx(data, filename="improved_resume.docx"):
    doc = docx.Document()
    
    name_p = doc.add_heading(data.get('name', 'Your Name'), level=1)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p = doc.add_paragraph(data.get('contact', 'Email | Phone'))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section_header(title): doc.add_heading(title, level=2)
        
    if data.get('summary'):
        add_section_header("Professional Summary")
        doc.add_paragraph(data['summary'])
        
    if data.get('skills') and len(data['skills']) > 0:
        add_section_header("Skills")
        for skill_group in data['skills']:
            p = doc.add_paragraph()
            cat = skill_group.get('category', '')
            items = skill_group.get('items', '')
            p.paragraph_format.left_indent = docx.shared.Inches(0.25)
            if cat: p.add_run(f"{cat}: ").bold = True
            p.add_run(items)

    # --- UPDATED EXPERIENCE DOCX ---
    if data.get('experience'):
        add_section_header("Experience")
        for item in data['experience']:
            p = doc.add_paragraph()
            p.add_run(item.get('title', '')).bold = True
            if item.get('date'): p.add_run(f"  |  {item.get('date', '')}")
            
            sub_parts = []
            if item.get('company'): sub_parts.append(item['company'])
            if item.get('technologies'): sub_parts.append(item['technologies'])
            sub_str = " | ".join(sub_parts)
            if item.get('location'): sub_str += f"  |  {item['location']}"
            
            if sub_str:
                p_sub = doc.add_paragraph()
                p_sub.add_run(sub_str).italic = True
                
            if item.get('details'):
                for detail in item['details']:
                    if str(detail).strip(): doc.add_paragraph(f"•  {str(detail)}")

    # --- UPDATED PROJECTS DOCX ---
    if data.get('projects'):
        add_section_header("Projects")
        for item in data['projects']:
            p = doc.add_paragraph()
            p.add_run(item.get('title', '')).bold = True
            if item.get('technologies'): p.add_run(f" | {item['technologies']}").italic = True
            if item.get('date'): p.add_run(f"  |  {item.get('date', '')}")
            
            if item.get('details'):
                for detail in item['details']:
                    if str(detail).strip(): doc.add_paragraph(f"•  {str(detail)}")
    
    # Education
    if data.get('education'):
        add_section_header("Education")
        for item in data['education']:
            p = doc.add_paragraph()
            p.add_run(item.get('title', '')).bold = True
            if item.get('date'): p.add_run(f"  |  {item.get('date', '')}")
            if item.get('subtitle'):
                p_sub = doc.add_paragraph()
                p_sub.add_run(item.get('subtitle', '')).italic = True
            if item.get('details'):
                for detail in item['details']:
                    if str(detail).strip(): doc.add_paragraph(f"•  {str(detail)}")

    for section_title, key in [("Certifications", "certifications"), ("Achievements & Leadership", "achievements")]:
        items = data.get(key, [])
        if items and len(items) > 0:
            add_section_header(section_title)
            for item in items:
                if str(item).strip(): doc.add_paragraph(f"•  {str(item)}")
                
    doc.save(filename)
    return filename

@app.route('/upload', methods=['POST'])
def upload():
    try:
        if 'resume' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
            
        file = request.files['resume']
        file_path = "temp_" + file.filename
        file.save(file_path)

        text = ""
        if file.filename.lower().endswith(".pdf"): text = extract_pdf(file_path)
        elif file.filename.lower().endswith(".docx"): text = extract_docx(file_path)
        else: return jsonify({"error": "Unsupported format"}), 400

        structured_data = enhance_with_ai(text)
        
        if "error" in structured_data:
             return jsonify(structured_data), 500

        pdf_file = generate_pdf(structured_data)
        docx_file = generate_docx(structured_data)
        
        if os.path.exists(file_path): os.remove(file_path)

        return jsonify({
            "resumeData": structured_data,
            "downloadUrl": f"http://127.0.0.1:5000/download/{pdf_file}"
        })

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    if not os.path.exists(filename):
        return "File not found! Please click 'Upload & Enhance' to generate this document again.", 404
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000)